from __future__ import annotations
import io
import logging
from typing import Callable

logger = logging.getLogger(__name__)

# ── PDF ───────────────────────────────────────────────────────────────────────

def extract_pdf(data: bytes) -> str:
    """Extract text and tables from PDF files."""
    import pdfplumber
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            table_text = ""
            for tbl in tables:
                rows = ["\t".join(str(c) if c else "" for c in row) for row in tbl if row]
                table_text += "\n".join(rows) + "\n"
            pages.append(f"--- Page {i} ---\n{text}\n{table_text}".strip())
    return "\n\n".join(pages)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def extract_docx(data: bytes) -> str:
    """Extract text from DOCX files."""
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for i, table in enumerate(doc.tables, 1):
        parts.append(f"[Table {i}]")
        for row in table.rows:
            parts.append("\t".join(cell.text.strip() for cell in row.cells))

    return "\n".join(parts)


# ── DOC (legacy) ──────────────────────────────────────────────────────────────

def extract_doc(data: bytes) -> str:
    """Uses antiword via subprocess; falls back to raw string extraction."""
    import subprocess, tempfile, os
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True, text=True, timeout=30, check=False
        )
        if result.returncode == 0:
            return result.stdout
        logger.warning("antiword failed, falling back to raw extraction")
    except FileNotFoundError:
        logger.warning("antiword not installed, using raw extraction")
    finally:
        os.unlink(tmp_path)

    # Raw fallback — strips most binary noise
    text = data.decode("latin-1", errors="ignore")
    printable = "".join(c if c.isprintable() or c in "\n\r\t" else " " for c in text)
    return " ".join(printable.split())


# ── XLSX ──────────────────────────────────────────────────────────────────────

def extract_xlsx(data: bytes) -> str:
    """Extract text from XLSX files."""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    sheets: list[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append("\t".join("" if v is None else str(v) for v in row))
        if rows:
            sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(sheets)

# ── XLS (legacy) ──────────────────────────────────────────────────────────────

def extract_xls(data: bytes) -> str:
    """Extract text from XLS files."""
    import xlrd
    wb = xlrd.open_workbook(file_contents=data)
    sheets: list[str] = []
    for sheet_name in wb.sheet_names():
        ws = wb.sheet_by_name(sheet_name)
        rows: list[str] = []
        for rx in range(ws.nrows):
            row = ws.row(rx)
            rows.append("\t".join(str(cell.value) for cell in row))
        if rows:
            sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
    return "\n\n".join(sheets)


# ── Images (OCR) ──────────────────────────────────────────────────────────────

def extract_image(data: bytes) -> str:
    """Extract text from image files using OCR."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img)
        return text.strip() or "[No readable text detected in image]"
    except ImportError:
        return "[OCR unavailable: install pytesseract + Tesseract]"
    except (OSError, RuntimeError, ValueError) as e:
        return f"[OCR failed: {e}]"


# ── Plain text / CSV ──────────────────────────────────────────────────────────

def extract_text(data: bytes) -> str:
    """Extract text from plain text files."""
    import chardet
    detected = chardet.detect(data)
    encoding = detected.get("encoding") or "utf-8"
    try:
        return data.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        return data.decode("utf-8", errors="replace")


def extract_csv(data: bytes) -> str:
    """Extract text from CSV files."""
    import csv, chardet
    detected = chardet.detect(data)
    encoding = detected.get("encoding") or "utf-8"
    text = data.decode(encoding, errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = ["\t".join(row) for row in reader]
    return "\n".join(rows)


# ── Router ────────────────────────────────────────────────────────────────────

_ROUTER: dict[str, Callable[[bytes], str]] = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "doc": extract_doc,
    "xlsx": extract_xlsx,
    "xls": extract_xls,
    "jpg": extract_image,
    "png": extract_image,
    "gif": extract_image,
    "tiff": extract_image,
    "bmp": extract_image,
    "txt": extract_text,
    "csv": extract_csv,
}


def extract_text_from_bytes(fmt: str, data: bytes) -> str:
    """
    Dispatch to the correct extractor.
    Returns extracted text or an error string.
    """
    fn = _ROUTER.get(fmt)
    if fn is None:
        return f"[Unsupported format: {fmt}]"
    try:
        return fn(data)
    except (OSError, RuntimeError, ValueError) as e:
        logger.error("Extraction failed for format '%s': %s", fmt, e, exc_info=True)
        return f"[Extraction error ({fmt}): {e}]"
